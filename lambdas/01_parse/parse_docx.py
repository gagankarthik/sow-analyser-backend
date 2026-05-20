"""DOCX → paragraph + table text extraction via python-docx."""
from __future__ import annotations

import io
from typing import Any


def parse_docx_bytes(data: bytes) -> dict[str, Any]:
    """Extract text from a .docx blob.

    Returns
    -------
    {"text": str, "pages": [{page, text, char_count}, ...]}

    DOCX has no inherent page concept; we emit a single synthetic page so the
    schema is uniform with PDF output.  If we ever need real page splitting
    we can render to PDF first, but for clause classification that's overkill.
    """
    doc = _load(data)
    paragraphs: list[str] = []

    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            paragraphs.append(t)

    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            cells = [c for c in cells if c]
            if cells:
                paragraphs.append(" | ".join(cells))

    full = "\n".join(paragraphs)
    return {
        "text": full,
        "pages": [
            {"page": 1, "text": full, "char_count": len(full)},
        ],
    }


def _load(data: bytes):
    from docx import Document  # type: ignore

    return Document(io.BytesIO(data))
