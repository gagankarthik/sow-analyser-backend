"""Stage 01 — Parse: extract raw text from DOCX or PDF.

Strategy:
  DOCX  → python-docx (direct)
  PDF   → pdfplumber first (fast, native text)
        → Textract async fallback (scanned / image-only PDFs)

Textract is only invoked when pdfplumber yields fewer than 200 total characters,
meaning the PDF is scanned/image-only. A single page with sparse text (cover,
blank, ToC) does NOT trigger Textract — only truly empty extractions do.
"""
from __future__ import annotations

import io
import time
from typing import Any

from shared.aws import textract_client
from shared.dynamodb import update_status
from shared.logger import get_logger
from shared.s3 import get_object, processed_key, put_json
from shared.schema import ExtractionMethod, now_iso
from shared.text import sha256_hex

log = get_logger("blue-iq.parse")

_PDF_MAGIC  = b"%PDF-"
_DOCX_MAGIC = b"PK\x03\x04"

TEXTRACT_POLL_INTERVAL_S = 5
TEXTRACT_MAX_WAIT_S      = 240


# ---------------------------------------------------------------------------
# Stage entry point
# ---------------------------------------------------------------------------


def run(event: dict[str, Any]) -> dict[str, Any]:
    raw_bucket       = event["rawBucket"]
    raw_key          = event["rawKey"]
    processed_bucket = event["processedBucket"]

    # Parse tenantId / docId from key: tenants/<tenantId>/uploads/<docId>/<file>
    doc_id, tenant_id = _ids_from_key(raw_key, event)
    event["docId"]    = doc_id
    event["tenantId"] = tenant_id

    log.append_keys(docId=doc_id, tenantId=tenant_id)
    log.info("parse.start", rawKey=raw_key)

    try:
        update_status(doc_id, "PARSING")
    except Exception:
        pass  # Row may not exist yet on very early failures; non-fatal.

    blob     = get_object(raw_bucket, raw_key)
    checksum = sha256_hex(blob)
    ftype    = _detect_type(raw_key, blob)

    if ftype == "docx":
        extracted = _parse_docx(blob)
        method    = ExtractionMethod.DOCX
    else:
        extracted = _try_pdfplumber(blob)
        if extracted and _has_content(extracted["pages"]):
            method = ExtractionMethod.PDFPLUMBER
        else:
            log.info("parse.pdf.fallback_to_textract", reason="insufficient pdfplumber text")
            extracted = _textract_async(raw_bucket, raw_key)
            method    = ExtractionMethod.TEXTRACT

    parsed = {
        "text":              extracted["text"],
        "pages":             extracted["pages"],
        "extracted_at":      now_iso(),
        "extraction_method": method.value,
        "checksum":          checksum,
    }
    out_key = processed_key(tenant_id, doc_id, "parsed.json")
    put_json(processed_bucket, out_key, parsed)
    log.info("parse.done", method=method.value, pages=len(parsed["pages"]),
             chars=len(parsed["text"]))

    event["parsed"] = parsed
    return event


# ---------------------------------------------------------------------------
# DOCX extraction
# ---------------------------------------------------------------------------


def _parse_docx(data: bytes) -> dict[str, Any]:
    from docx import Document  # type: ignore

    doc        = Document(io.BytesIO(data))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    full = "\n".join(paragraphs)
    return {"text": full, "pages": [{"page": 1, "text": full, "char_count": len(full)}]}


# ---------------------------------------------------------------------------
# PDF — pdfplumber (native text layer)
# ---------------------------------------------------------------------------


def _try_pdfplumber(data: bytes) -> dict[str, Any] | None:
    try:
        import pdfplumber  # type: ignore
    except ImportError:
        return None

    pages: list[dict[str, Any]] = []
    try:
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for i, page in enumerate(pdf.pages, start=1):
                text = (page.extract_text() or "").strip()
                pages.append({"page": i, "text": text, "char_count": len(text)})
    except Exception as exc:
        log.warning("parse.pdf.pdfplumber_error", error=str(exc))
        return None

    full = "\n\n".join(p["text"] for p in pages)
    return {"text": full, "pages": pages}


def _has_content(pages: list[dict[str, Any]]) -> bool:
    """Return True when the document has enough text to be usable.

    Only the TOTAL character count matters — a single sparse page (cover,
    blank, table-of-contents) should not force an expensive Textract call.
    """
    return sum(p["char_count"] for p in pages) >= 200


# ---------------------------------------------------------------------------
# PDF — Textract async (scanned / image-only)
# ---------------------------------------------------------------------------


def _textract_async(bucket: str, key: str) -> dict[str, Any]:
    tx     = textract_client()
    job_id = tx.start_document_text_detection(
        DocumentLocation={"S3Object": {"Bucket": bucket, "Name": key}}
    )["JobId"]
    log.info("textract.started", jobId=job_id)

    waited = 0
    while True:
        if waited >= TEXTRACT_MAX_WAIT_S:
            raise TimeoutError(f"Textract job {job_id} timed out after {TEXTRACT_MAX_WAIT_S}s")
        time.sleep(TEXTRACT_POLL_INTERVAL_S)
        waited += TEXTRACT_POLL_INTERVAL_S
        status  = tx.get_document_text_detection(JobId=job_id, MaxResults=1)["JobStatus"]
        if status != "IN_PROGRESS":
            break

    if status != "SUCCEEDED":
        raise RuntimeError(f"Textract job {job_id} failed with status: {status}")

    blocks: list[dict[str, Any]] = []
    next_token: str | None = None
    while True:
        kwargs: dict[str, Any] = {"JobId": job_id, "MaxResults": 1000}
        if next_token:
            kwargs["NextToken"] = next_token
        resp       = tx.get_document_text_detection(**kwargs)
        blocks    += resp.get("Blocks", [])
        next_token = resp.get("NextToken")
        if not next_token:
            break

    page_lines: dict[int, list[str]] = {}
    for b in blocks:
        if b.get("BlockType") == "LINE":
            page_lines.setdefault(int(b.get("Page", 1)), []).append(b.get("Text", ""))

    page_list = [
        {"page": p, "text": "\n".join(lines), "char_count": sum(len(l) for l in lines)}
        for p, lines in sorted(page_lines.items())
    ]
    full = "\n\n".join(p["text"] for p in page_list)
    log.info("textract.done", pages=len(page_list), chars=len(full))
    return {"text": full, "pages": page_list}


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _ids_from_key(raw_key: str, event: dict[str, Any]) -> tuple[str, str]:
    """Extract (docId, tenantId) from key: tenants/<tenantId>/uploads/<docId>/<file>."""
    parts = raw_key.split("/")
    if len(parts) >= 4 and parts[0] == "tenants" and parts[2] == "uploads":
        return parts[3], parts[1]
    return event.get("docId", raw_key), event.get("tenantId", "default")


def _detect_type(filename: str, blob: bytes) -> str:
    name = filename.lower()
    head = blob[:8]
    if name.endswith(".pdf") or head.startswith(_PDF_MAGIC) or _PDF_MAGIC in blob[:1024]:
        return "pdf"
    if name.endswith(".docx") or head.startswith(_DOCX_MAGIC):
        return "docx"
    raise ValueError(f"Unsupported file type: {filename!r}. Supported: pdf, docx")
